"""Generate SDIE V6 model flow DOCX."""
from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    raise SystemExit("Install python-docx: pip install python-docx")

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SDIE_V6_Model_Flow.docx"


def _shade_cell(cell, hex_fill: str) -> None:
    shading = cell._tc.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): hex_fill, qn("w:val"): "clear"})
    shading.append(shd)


def main() -> None:
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_heading("SDIE V6 — Model Flow", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Structural Drawing Intelligence Engine: teach from reference projects "
        "(INIZIO, TRUST_OFFICE, MANOHAR), then infer and quantify on any raw DXF "
        "using GENERIC merged knowledge. Outputs include slab and beam concrete takeoff."
    )

    # --- Section 1: End-to-end ---
    doc.add_heading("1. End-to-End Flow", 1)
    doc.add_paragraph(
        "Two phases: (A) offline teach from tagged reference DXFs, "
        "(B) runtime inference on any drawing."
    )

    flow1 = """
    ┌─────────────────────────────────────────────────────────────────┐
    │  A · TEACH (offline)                                            │
    │  Tagged DXFs → Atlas + Layer Profiles + Knowledge Base          │
    └────────────────────────────┬────────────────────────────────────┘
                                 │ merged artifacts
                                 ▼
    ┌─────────────────────────────────────────────────────────────────┐
    │  B · INFER (runtime, project-id GENERIC)                        │
    │                                                                 │
    │  Input DXF                                                      │
    │      → Prepare (layers, floor zone, entities)                   │
    │      → Classify (rules, graph, DeepSeek)                        │
    │      → Quantify (beams + slabs)                                 │
    │      → Export (JSON, overlay, Excel, review queue)            │
    └─────────────────────────────────────────────────────────────────┘
    """
    p = doc.add_paragraph()
    r = p.add_run(flow1.strip())
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    # --- Section 2: Teach ---
    doc.add_heading("2. Teach Pipeline (Phase A)", 1)
    teach_steps = [
        "Read projects_manifest.json",
        "Load tagged DXF per reference project",
        "build_atlas.py → component_atlas.json",
        "build_layer_profiles.py → layer_profiles.json",
        "RAG / KB builder → structural_kb.json",
    ]
    for i, step in enumerate(teach_steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")

    # --- Section 3: Inference steps ---
    doc.add_heading("3. Inference Pipeline (Phase B)", 1)
    infer_table = doc.add_table(rows=1, cols=2)
    infer_table.style = "Table Grid"
    infer_table.rows[0].cells[0].text = "Step"
    infer_table.rows[0].cells[1].text = "Action"
    _shade_cell(infer_table.rows[0].cells[0], "E8F4FD")
    _shade_cell(infer_table.rows[0].cells[1], "E8F4FD")

    steps = [
        ("①", "Load DXF — units, scale, extents"),
        ("②", "Auto-discover layers — frame, annotation, void, column"),
        ("③", "Resolve floor zone — THK label clusters"),
        ("④", "Load teach artifacts — Atlas + KB (GENERIC = all projects)"),
        ("⑤", "Extract entities — structural, cutout, hatch void, columns"),
        ("⑥", "Classify components — V5: rules → graph → DeepSeek on ambiguous"),
        ("⑦", "Beam quantities — centerline length × section → concrete m³"),
        ("⑧", "Build structural graph — nodes and topology edges"),
        ("⑨", "Slab detection — exclusions → beam grid → framed bays"),
        ("⑩", "Slab quantities — area × THK → concrete m³"),
        ("⑪", "Export — JSON, overlay, Excel (Slabs + Beams), review queue"),
    ]
    for num, action in steps:
        row = infer_table.add_row().cells
        row[0].text = num
        row[1].text = action

    # --- Section 4: Classification ---
    doc.add_heading("4. V5 Classification Logic", 1)
    for item in [
        "Hard global layers → locked type (Opening, Column, Beam, etc.)",
        "Soft layers (S_FRAMES, STR-BEAM) → geometry heuristics first",
        "GENERIC mode → best teach rule per layer across all projects",
        "V5 confidence < 75% → DeepSeek with RAG + graph context",
        "Supporting types excluded from slab area: walls, cores, shafts, openings",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # --- Section 5: Quantities ---
    doc.add_heading("5. Quantity Engines", 1)
    qty_table = doc.add_table(rows=1, cols=3)
    qty_table.style = "Table Grid"
    for i, h in enumerate(["Element", "Inputs", "Formula"]):
        qty_table.rows[0].cells[i].text = h
        _shade_cell(qty_table.rows[0].cells[i], "E6F4EA")

    for row_data in [
        ("Beam", "Classified beam lines, beam-size tags", "length × width × depth → m³"),
        ("Slab", "Framed bays after exclusions, THK labels", "area × thickness → m³"),
        ("Combined", "Slab + beam totals", "Reported in Summary and Excel"),
    ]:
        row = qty_table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val

    doc.add_paragraph("Slab fallback: beam-grid → region polygonize → beam-frame bbox.")

    # --- Section 6: Inference mode ---
    doc.add_heading("6. Inference Mode", 1)
    mode_table = doc.add_table(rows=1, cols=3)
    mode_table.style = "Table Grid"
    for i, h in enumerate(["project-id", "Knowledge", "Use case"]):
        mode_table.rows[0].cells[i].text = h
    for a, b, c in [
        ("GENERIC (default)", "All teach projects merged", "New / unknown drawings"),
        ("INIZIO", "Inizio + GLOBAL", "Teach evaluation"),
        ("TRUST_OFFICE", "Trust + GLOBAL", "Teach evaluation"),
    ]:
        row = mode_table.add_row().cells
        row[0].text, row[1].text, row[2].text = a, b, c

    # --- Section 7: CLI ---
    doc.add_heading("7. Key CLI Flags", 1)
    for item in [
        "python scripts/run_pipeline.py <dxf> -o <output_dir>",
        "--project-id GENERIC (default)",
        "--no-deepseek — skip DeepSeek reasoning",
        "--no-beam-quantities — skip beam takeoff",
        "DEEPSEEK_API_KEY in Phase2_Concrete_Estimation/.env",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # --- Section 8: Outputs ---
    doc.add_heading("8. Output Files", 1)
    for item in [
        "{drawing}_results.json",
        "{drawing}_building_model.json",
        "{drawing}_overlay.html / .svg",
        "{drawing}_quantities.xlsx — Summary, Slabs, Beams, Classification",
        "{drawing}_review_queue.json",
        "{drawing}_summary.txt",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_paragraph()
    foot = doc.add_paragraph("SDIE V6 · P2_SlabVersion6 · Teach-then-infer structural estimation")
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot.runs[0].font.size = Pt(9)
    foot.runs[0].font.italic = True
    foot.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
